"""Extract structured text from PDF / DOCX assignment files.

Returns a ``ParsedAssignment`` containing sections with headings and
individual paragraphs, preserving as much document structure as possible.
"""

from __future__ import annotations

import logging
import re
from io import BytesIO
from typing import List

from models import ParsedAssignment, ParsedSection

logger = logging.getLogger(__name__)

# Heuristics for detecting heading-like lines in raw PDF text.
_HEADING_MAX_WORDS = 12
_HEADING_PATTERN = re.compile(
    r"^(?:"
    r"(?:chapter|section|part)\s+\d+|"  # "Chapter 1", "Section 3"
    r"\d+(?:\.\d+)*\s+\S|"               # "1.2 Some title"
    r"[A-Z][A-Z\s]{3,}$"                 # ALL-CAPS lines
    r")",
    re.IGNORECASE,
)


def parse_assignment(filename: str, file_bytes: bytes) -> ParsedAssignment:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if suffix == "pdf":
        sections = _parse_pdf(file_bytes)
    elif suffix == "docx":
        sections = _parse_docx(file_bytes)
    else:
        raise ValueError("Unsupported assignment file type. Only PDF and DOCX are allowed.")

    full_text = _sections_to_text(sections)
    word_count = len(full_text.split())
    logger.info("Parsed assignment: %d sections, %d words", len(sections), word_count)
    return ParsedAssignment(sections=sections, full_text=full_text, word_count=word_count)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(file_bytes: bytes) -> List[ParsedSection]:
    from pypdf import PdfReader  # lazy import keeps startup fast

    reader = PdfReader(BytesIO(file_bytes))
    lines: List[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        for line in page_text.splitlines():
            stripped = line.strip()
            if stripped:
                lines.append(stripped)

    return _lines_to_sections(lines)


def _is_likely_heading(line: str) -> bool:
    if len(line.split()) > _HEADING_MAX_WORDS:
        return False
    if _HEADING_PATTERN.match(line):
        return True
    # Short line that is title-cased (and not just a single word)
    words = line.split()
    if 2 <= len(words) <= _HEADING_MAX_WORDS and line == line.title():
        return True
    return False


def _lines_to_sections(lines: List[str]) -> List[ParsedSection]:
    """Group raw lines into sections using heading heuristics."""
    sections: List[ParsedSection] = []
    current_title: str | None = None
    current_paragraphs: List[str] = []
    buffer: List[str] = []

    def flush_buffer() -> None:
        nonlocal buffer
        if buffer:
            current_paragraphs.append(" ".join(buffer))
            buffer = []

    for line in lines:
        if _is_likely_heading(line):
            flush_buffer()
            if current_paragraphs or current_title is not None:
                sections.append(ParsedSection(title=current_title, paragraphs=current_paragraphs))
            current_title = line
            current_paragraphs = []
        else:
            # Blank-line-separated paragraphs are already split by the caller
            # (each non-empty line is its own entry). Accumulate contiguous
            # lines into paragraphs and break on very short "transition" lines.
            if len(line.split()) <= 3 and buffer:
                flush_buffer()
            buffer.append(line)

    flush_buffer()
    if current_paragraphs or current_title is not None:
        sections.append(ParsedSection(title=current_title, paragraphs=current_paragraphs))
    elif not sections and current_paragraphs:
        sections.append(ParsedSection(title=None, paragraphs=current_paragraphs))

    # If no headings were detected at all, put everything in one section.
    if not sections and lines:
        sections.append(ParsedSection(title=None, paragraphs=[" ".join(lines)]))

    return sections


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _parse_docx(file_bytes: bytes) -> List[ParsedSection]:
    import docx  # lazy import

    document = docx.Document(BytesIO(file_bytes))
    sections: List[ParsedSection] = []
    current_title: str | None = None
    current_paragraphs: List[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower() if para.style else ""
        is_heading = "heading" in style_name or "title" in style_name

        if is_heading:
            if current_paragraphs or current_title is not None:
                sections.append(ParsedSection(title=current_title, paragraphs=current_paragraphs))
            current_title = text
            current_paragraphs = []
        else:
            current_paragraphs.append(text)

    if current_paragraphs or current_title is not None:
        sections.append(ParsedSection(title=current_title, paragraphs=current_paragraphs))

    if not sections:
        all_text = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        sections.append(ParsedSection(title=None, paragraphs=all_text))

    return sections


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _sections_to_text(sections: List[ParsedSection]) -> str:
    parts: List[str] = []
    for sec in sections:
        if sec.title:
            parts.append(sec.title)
        parts.extend(sec.paragraphs)
    return "\n\n".join(parts)
