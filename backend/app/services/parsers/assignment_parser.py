"""Extract structured text from PDF / DOCX assignment files."""

from __future__ import annotations

import logging
import re
from io import BytesIO
from typing import List

from app.services.analysis.models import ParsedAssignment, ParsedSection

logger = logging.getLogger(__name__)

_HEADING_MAX_WORDS = 12
_HEADING_PATTERN = re.compile(
    r"^(?:"
    r"(?:chapter|section|part)\s+\d+|"
    r"\d+(?:\.\d+)*\s+\S|"
    r"[A-Z][A-Z\s]{3,}$"
    r")",
    re.IGNORECASE,
)


def parse_assignment(filename: str, file_bytes: bytes) -> ParsedAssignment:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if suffix == "pdf":
        sections = _parse_pdf(file_bytes)
    elif suffix == "docx":
        sections = _parse_docx(file_bytes)
    else:
        raise ValueError("Unsupported assignment file type. Only PDF and DOCX are allowed.")

    full_text = _sections_to_text(sections)
    word_count = len(full_text.split())
    logger.info("Parsed assignment: %d sections, %d words", len(sections), word_count)
    return ParsedAssignment(sections=sections, full_text=full_text, word_count=word_count)


def _parse_pdf(file_bytes: bytes) -> List[ParsedSection]:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(file_bytes))
    lines: List[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        for line in page_text.splitlines():
            stripped = line.strip()
            if stripped:
                lines.append(stripped)

    return _lines_to_sections(lines)


def _is_likely_heading(line: str) -> bool:
    if len(line.split()) > _HEADING_MAX_WORDS:
        return False
    if _HEADING_PATTERN.match(line):
        return True
    words = line.split()
    if 2 <= len(words) <= _HEADING_MAX_WORDS and line == line.title():
        return True
    return False


def _lines_to_sections(lines: List[str]) -> List[ParsedSection]:
    sections: List[ParsedSection] = []
    current_title: str | None = None
    current_paragraphs: List[str] = []
    buffer: List[str] = []

    def flush_buffer() -> None:
        nonlocal buffer
        if buffer:
            current_paragraphs.append(" ".join(buffer))
            buffer = []

    for line in lines:
        if _is_likely_heading(line):
            flush_buffer()
            if current_paragraphs or current_title is not None:
                sections.append(ParsedSection(title=current_title, paragraphs=current_paragraphs))
            current_title = line
            current_paragraphs = []
        else:
            if len(line.split()) <= 3 and buffer:
                flush_buffer()
            buffer.append(line)

    flush_buffer()
    if current_paragraphs or current_title is not None:
        sections.append(ParsedSection(title=current_title, paragraphs=current_paragraphs))
    elif not sections and current_paragraphs:
        sections.append(ParsedSection(title=None, paragraphs=current_paragraphs))

    if not sections and lines:
        sections.append(ParsedSection(title=None, paragraphs=[" ".join(lines)]))

    return sections


def _parse_docx(file_bytes: bytes) -> List[ParsedSection]:
    import docx

    document = docx.Document(BytesIO(file_bytes))
    sections: List[ParsedSection] = []
    current_title: str | None = None
    current_paragraphs: List[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower() if para.style else ""
        is_heading = "heading" in style_name or "title" in style_name

        if is_heading:
            if current_paragraphs or current_title is not None:
                sections.append(ParsedSection(title=current_title, paragraphs=current_paragraphs))
            current_title = text
            current_paragraphs = []
        else:
            current_paragraphs.append(text)

    if current_paragraphs or current_title is not None:
        sections.append(ParsedSection(title=current_title, paragraphs=current_paragraphs))

    if not sections:
        all_text = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        sections.append(ParsedSection(title=None, paragraphs=all_text))

    return sections


def _sections_to_text(sections: List[ParsedSection]) -> str:
    parts: List[str] = []
    for sec in sections:
        if sec.title:
            parts.append(sec.title)
        parts.extend(sec.paragraphs)
    return "\n\n".join(parts)
